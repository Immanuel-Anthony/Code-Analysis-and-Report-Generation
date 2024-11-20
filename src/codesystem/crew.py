from crewai import Agent, Task, Crew
import yaml
from pathlib import Path
from datetime import datetime
import os
import re  
from docx import Document  # Import python-docx for DOCX generation

class Codesystem:
    def __init__(self):
        self.config_dir = Path(__file__).parent / "config"
        self.output_dir = Path(__file__).parent / "reports"
        self.agents_config = self._load_yaml("agents.yaml")
        self.tasks_config = self._load_yaml("tasks.yaml")
        
        # Create output directory if it doesn't exist
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _load_yaml(self, filename):
        file_path = self.config_dir / filename
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return yaml.safe_load(file)
        except Exception as e:
            raise Exception(f"Error loading {filename}: {e}")

    def save_report(self, content, filename_base, file_name=None):
        """Save the content dynamically into a DOCX file."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Generate a meaningful filename from the input file
        if file_name:
            # Extract just the filename without the full path and extension
            base_name = Path(file_name).stem
            # Clean the filename to remove any invalid characters
            base_name = re.sub(r'[^\w\-_]', '_', base_name)
            docx_path = self.output_dir / f"{base_name}_analysis_{timestamp}.docx"
        else:
            docx_path = self.output_dir / f"{filename_base}_{timestamp}.docx"
        
        # Save the dynamic content to a DOCX file
        document = Document()

        # Add content to DOCX with proper formatting
        document.add_heading('Code Analysis Report', 0)
        
        # Split content into sections if it contains headers
        sections = content.split('\n# ')
        
        # Add the first section (or all content if no sections)
        document.add_paragraph(sections[0])
        
        # Add remaining sections with proper headers
        for section in sections[1:]:
            if section.strip():
                # Split section into title and content
                lines = section.split('\n', 1)
                if len(lines) > 0:
                    document.add_heading(lines[0], level=1)
                    if len(lines) > 1:
                        document.add_paragraph(lines[1])

        # Save the DOCX report
        document.save(str(docx_path))
        
        print(f"\nDOCX report saved as: {docx_path}")
        return docx_path

    def crew(self):
        # Create agents dynamically
        agents = {}
        for agent_name, agent_config in self.agents_config.items():
            agents[agent_name] = Agent(
                role=agent_config['role'],
                goal=agent_config['goal'],
                backstory=agent_config['backstory'],
                verbose=True
            )

        # Create tasks dynamically
        tasks = []
        task_outputs = {}

        # Modified task creation to ensure proper dependency chain
        for task_name, task_config in self.tasks_config.items():
            dependencies = []
            if 'dependencies' in task_config:
                for dep in task_config['dependencies']:
                    if dep in task_outputs:
                        dependencies.append(task_outputs[dep])

            task = Task(
                description=task_config['description'],
                expected_output=task_config.get('expected_output', None),
                agent=agents[task_config['agent']],
                dependencies=dependencies,
                verbose=True
            )
            tasks.append(task)
            task_outputs[task_name] = task

        # Create crew
        crew = Crew(
            agents=list(agents.values()),
            tasks=tasks,
            verbose=True
        )

        return crew, task_outputs

    def kickoff(self, inputs=None):
        """Start the crew execution with the given inputs"""
        if inputs is None:
            inputs = {}
        
        # Initialize all expected outputs to avoid KeyError
        initial_inputs = {
            'code_to_analyze': inputs.get('code_to_analyze', ''),
            'file_name': inputs.get('file_name', ''),
            'code_analysis_output': '',
            'security_analysis_output': '',
            'performance_analysis_output': '',
            'code_test_output': '',
            'best_practices_output': '',
            'final_report_task': ''
        }

        crew, task_outputs = self.crew()
        
        # Run the analysis and get the result
        result = crew.kickoff(inputs=initial_inputs)
        
        # Get the file name from inputs
        file_name = inputs.get('file_name', '') if inputs else ''
        
        # Ensure we have a result and save it to DOCX
        if result:
            if not isinstance(result, str):
                result = str(result)
            
            # Pass the file_name to save_report
            filename_base = "code_analysis_report"
            docx_path = self.save_report(result, filename_base, file_name=file_name)
            print(f"\nReport has been generated and saved to:\n{docx_path}")
        else:
            print("No result was generated to save to DOCX")
        
        return result

