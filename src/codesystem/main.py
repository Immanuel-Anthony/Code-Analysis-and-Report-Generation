#!/usr/bin/env python
import sys
import io
import re
from pathlib import Path
from crew import Codesystem  # Ensure this import points to the correct path for Codesystem
from docx import Document  # Import the python-docx library
import git
from pathlib import Path
import tempfile
import os

def clone_github_repo(github_url):
    """
    Clone a GitHub repository and return the local path.
    """
    # Create a temporary directory
    temp_dir = tempfile.mkdtemp()
    try:
        # Clone the repository
        git.Repo.clone_from(github_url, temp_dir)
        return temp_dir
    except Exception as e:
        raise Exception(f"Error cloning repository: {e}")

def get_code_files(repo_path):
    """
    Recursively get all code files from the repository.
    """
    code_extensions = {'.py', '.js', '.java', '.cpp', '.c', '.h', '.hpp', '.cs', '.rb', '.go', '.rs', '.php'}
    code_files = []
    
    for root, _, files in os.walk(repo_path):
        for file in files:
            file_path = Path(root) / file
            if file_path.suffix in code_extensions:
                code_files.append(file_path)
    
    return code_files

def read_code_file(file_path):
    """
    Read code from a file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Warning: Could not read {file_path}: {e}")
        return None

def prepare_inputs(file_path, code_content, agent_outputs):
    """
    Prepare inputs for report generation with the real outputs.
    """
    return {
        'code_to_analyze': code_content,
        'file_name': str(file_path),
        'code_analysis_output': agent_outputs.get('code_analysis_output', 'No output available'),
        'security_analysis_output': agent_outputs.get('security_analysis_output', 'No output available'),
        'performance_analysis_output': agent_outputs.get('performance_analysis_output', 'No output available'),
        'code_test_coverage': agent_outputs.get('code_test_output', 'No output available'),
        'best_practices_output': agent_outputs.get('best_practices_output', 'No output available')
    }

def clean_text(text):
    """
    Cleans text by replacing escape sequences with appropriate formatting.
    """
    if not isinstance(text, str):
        text = str(text)  # Convert non-string input to string
    text = re.sub(r'\\n', ' ', text)  # Replace literal '\\n' with a space
    text = re.sub(r'\\t', '    ', text)  # Replace '\\t' with a tab space
    text = re.sub(r'\n+', '\n', text)  # Replace multiple newlines with one newline
    text = text.replace('\r', '')  # Remove any carriage returns that may be in the text

    # Strip leading/trailing spaces and return cleaned text
    return text.strip()


def generate_docx_report(inputs, docx_file_path):
    """
    Generate the DOCX report from the inputs.
    """
    # Create the DOCX file
    doc = Document()

    # Add title to the document
    doc.add_heading('Generated Report', 0)

    # Add the file name
    doc.add_paragraph(f"File: {inputs.get('file_name')}")

    # Add the different sections with cleaned text
    doc.add_heading('Code Analysis Output', level=1)
    code_analysis = inputs.get('code_analysis_output', 'No output available')
    if hasattr(code_analysis, 'result'):  # If it's a Task object
        code_analysis = code_analysis.result()  # Get the actual result
    doc.add_paragraph(clean_text(code_analysis))

    doc.add_heading('Security Analysis Output', level=1)
    doc.add_paragraph(clean_text(inputs.get('security_analysis_output', 'No output available')))

    doc.add_heading('Performance Analysis Output', level=1)
    doc.add_paragraph(clean_text(inputs.get('performance_analysis_output', 'No output available')))

    doc.add_heading('Architecture Analysis Output', level=1)
    doc.add_paragraph(clean_text(inputs.get('code_test_output', 'No output available')))

    doc.add_heading('Best Practices Output', level=1)
    doc.add_paragraph(clean_text(inputs.get('best_practices_output', 'No output available')))

    # Save the DOCX file
    doc.save(docx_file_path)


def run(github_url=None, file_path=None):
    repo_path = None
    try:
        if github_url is None and file_path is None:
            raise ValueError("Please provide either a GitHub URL or a path to the code file")
        
        codesystem = Codesystem()
        
        if github_url:
            # Clone the repository
            repo_path = clone_github_repo(github_url)
            
            # Get all code files
            code_files = get_code_files(repo_path)
            
            # Process each file
            for file_path in code_files:
                print(f"\nProcessing: {file_path}")
                code_content = read_code_file(file_path)
                
                if code_content:
                    initial_inputs = {
                        'code_to_analyze': code_content,
                        'file_name': str(file_path)
                    }
                    
                    # Process the file
                    result = codesystem.kickoff(inputs=initial_inputs)
            
            return result
        else:
            # Original single file processing
            file_path = Path(file_path)
            code_content = read_code_file(file_path)
            initial_inputs = {
                'code_to_analyze': code_content,
                'file_name': str(file_path)
            }
            return codesystem.kickoff(inputs=initial_inputs)
            
    finally:
        # Cleanup temporary directory with error handling
        if repo_path and os.path.exists(repo_path):
            import shutil
            try:
                shutil.rmtree(repo_path, ignore_errors=True)
            except Exception as e:
                print(f"Warning: Could not fully remove temporary directory {repo_path}. You may need to remove it manually.")
                print(f"Error details: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python main.py <path_to_code_file_or_github_url>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    if input_path.startswith(('http://', 'https://', 'git://')):
        run(github_url=input_path)
    else:
        run(file_path=input_path)
